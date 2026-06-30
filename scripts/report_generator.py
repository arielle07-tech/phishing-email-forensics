#!/usr/bin/env python3
"""
Phishing Email Forensics — Report Generator
=============================================
Génère des rapports d'incident SOC en PDF et DOCX.

Structure alignée sur le framework SOC (clevertailor) :
 1. Executive Summary — what happened, impact, key takeaway
 2. Alert Details — source, time, severity, description
 3. Triage Decision — true/false positive and why
 4. Enrichment Findings — IP/Domain reputation, context, evidence
 5. MITRE ATT&CK Mapping — tactic, technique, ID, description
 6. Recommended Response — next steps and ownership
 A. Annexes techniques (headers bruts, analyse IA détaillée)
"""

import io
import os
import re
import hashlib
from datetime import datetime


# ════════════════════════════════════════
# Utility functions
# ════════════════════════════════════════

def _safe(text: str) -> str:
    if not text:
        return ""
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _format_bytes(b: int) -> str:
    if not b:
        return "0 B"
    for unit in ['B', 'KB', 'MB', 'GB']:
        if b < 1024:
            return f"{b:.1f} {unit}"
        b /= 1024
    return f"{b:.1f} TB"


def _defang_url(url: str) -> str:
    if not url:
        return ""
    url = url.replace("http://", "hxxp://").replace("https://", "hxxps://")
    url = url.replace("://", "[://]")
    parts = url.split("/", 3)
    if len(parts) >= 3:
        parts[2] = parts[2].replace(".", "[.]")
        return "/".join(parts)
    return url.replace(".", "[.]")


def _defang_ip(ip: str) -> str:
    if not ip:
        return ""
    return ip.replace(".", "[.]")


def _defang_domain(domain: str) -> str:
    if not domain:
        return ""
    return domain.replace(".", "[.]")


def _generate_incident_id(analysis: dict) -> str:
    meta = analysis.get("metadata", {})
    seed = f"{meta.get('message_id', '')}{meta.get('date', '')}{meta.get('from', '')}"
    h = hashlib.md5(seed.encode()).hexdigest()[:6].upper()
    return f"INC-{datetime.now().strftime('%Y')}-{h}"


def _generate_report_number(analysis: dict) -> str:
    """Génère un numéro de rapport unique pour suivi (RPT-YYYYMMDD-XXXXXX)."""
    meta = analysis.get("metadata", {})
    seed = f"{meta.get('message_id', '')}{datetime.now().isoformat()}{meta.get('from', '')}"
    h = hashlib.md5(seed.encode()).hexdigest()[:6].upper()
    return f"RPT-{datetime.now().strftime('%Y%m%d')}-{h}"


def _clean_recommendation(text: str) -> str:
    """Supprime la numérotation en début de recommandation (ex: '1. ', '2. ')."""
    import re
    return re.sub(r'^\d+[\.\)\-]\s*', '', text.strip())


def _classify_incident(analysis: dict) -> dict:
    ai = analysis.get("ai_analysis", {})
    iocs = analysis.get("iocs", {})
    atts = analysis.get("attachments", [])
    verdict = analysis.get("verdict", {})
    score = verdict.get("score", 0)

    has_malicious_att = any(a.get("suspicious_extension") for a in atts)
    has_credential_keywords = any(
        kw in (iocs.get("phishing_keywords", []))
        for kw in ["verify your account", "password expired", "reset your password",
                    "confirm your identity", "vérifiez votre compte", "mot de passe expiré"]
    )
    imp_entity = ai.get("impersonation", {}).get("impersonated_entity", "")

    if has_malicious_att:
        category = "Malware Delivery"
        sub_category = "Phishing avec pièce jointe malveillante"
    elif has_credential_keywords:
        category = "Credential Harvesting"
        sub_category = "Vol d'identifiants via formulaire frauduleux"
    elif imp_entity and imp_entity != "N/A":
        category = "Business Email Compromise (BEC)"
        sub_category = f"Usurpation d'identité — {imp_entity}"
    else:
        category = "Phishing générique"
        sub_category = "Tentative de phishing non catégorisée"

    level = verdict.get("risk_level", "INFO")
    severity_map = {
        "CRITICAL": ("P1", "Critique — Réponse immédiate requise"),
        "HIGH": ("P2", "Élevée — Traitement prioritaire"),
        "MEDIUM": ("P3", "Modérée — Traitement dans les 24h"),
        "LOW": ("P4", "Faible — Surveillance recommandée"),
        "INFO": ("P5", "Informationnel — Aucune action requise"),
    }
    priority, severity_desc = severity_map.get(level, ("P5", "Non classifié"))

    return {
        "category": category,
        "sub_category": sub_category,
        "priority": priority,
        "severity_description": severity_desc,
        "tlp": "TLP:AMBER" if score >= 45 else "TLP:GREEN",
        "mitre_tactics": [m.get("id", "") for m in ai.get("mitre_techniques", [])],
    }


def _build_executive_summary(analysis: dict, classification: dict, incident_id: str) -> str:
    """Construit un résumé exécutif one-liner + contexte, même sans IA."""
    verdict = analysis.get("verdict", {})
    meta = analysis.get("metadata", {})
    ai = analysis.get("ai_analysis", {})
    score = verdict.get("score", 0)
    level = verdict.get("risk_level", "INFO")

    # Use AI executive summary if available
    ai_summary = ""
    if ai and not ai.get("error"):
        ai_summary = ai.get("executive_summary", "")

    # Build the one-liner (always present)
    sender = meta.get("from", "Inconnu")
    subject = meta.get("subject", "Sans objet")
    one_liner = (
        f"Severity: {level} — {classification['category']} "
        f"detecte depuis {_safe(sender)} "
        f"(score {score}/100)."
    )

    if ai_summary:
        return f"{one_liner}\n\n{ai_summary}"
    else:
        verdict_summary = verdict.get("summary", "")
        if verdict_summary:
            return f"{one_liner}\n\n{verdict_summary}"
        return one_liner


def _build_timeline(analysis: dict) -> list:
    meta = analysis.get("metadata", {})
    timeline = []
    email_date = meta.get("date", "")
    analysis_ts = meta.get("analysis_timestamp", "")
    now = datetime.now().strftime("%d/%m/%Y %H:%M UTC")

    if email_date:
        timeline.append(("Reception", email_date, "Email suspect recu par le destinataire"))
    timeline.append(("Soumission", analysis_ts or now, "Email soumis a la plateforme d'analyse"))
    timeline.append(("Analyse", now, "Extraction IOCs, authentification, scoring"))
    ai = analysis.get("ai_analysis", {})
    if ai and not ai.get("error"):
        timeline.append(("Analyse IA", now, "Tactiques social engineering identifiees"))
    timeline.append(("Rapport", now, "Rapport genere, pret pour escalade"))
    return timeline


# ════════════════════════════════════════════════════════════════
#
#   PDF REPORT
#
# ════════════════════════════════════════════════════════════════

def generate_pdf_report(analysis: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, KeepTogether
    )

    buf = io.BytesIO()

    # ── Page dimensions ──
    PAGE_W = A4[0]
    MARGIN_L = 22*mm
    MARGIN_R = 22*mm
    CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R  # ~166mm

    # Extract data early (needed for footer)
    meta = analysis.get("metadata", {})
    verdict = analysis.get("verdict", {})
    auth = analysis.get("authentication", {})
    iocs = analysis.get("iocs", {})
    atts = analysis.get("attachments", [])
    risk = analysis.get("risk_scoring", {})
    ai = analysis.get("ai_analysis", {})
    headers_a = analysis.get("headers_analysis", {})

    score = verdict.get("score", 0)
    level = verdict.get("risk_level", "INFO")
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    incident_id = _generate_incident_id(analysis)
    report_number = _generate_report_number(analysis)
    classification = _classify_incident(analysis)

    # ── Footer on every page ──
    def _page_footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont('Helvetica', 7)
        canvas.setFillColor(HexColor('#a0aec0'))
        footer_text = (
            f"Rapport {report_number}  |  Incident {incident_id}  |  "
            f"Phishing Email Forensics  |  {classification['tlp']}  |  CONFIDENTIEL"
        )
        canvas.drawCentredString(PAGE_W / 2, 12*mm, footer_text)
        canvas.drawRightString(PAGE_W - MARGIN_R, 12*mm, f"Page {doc_obj.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=20*mm, bottomMargin=22*mm,
                            leftMargin=MARGIN_L, rightMargin=MARGIN_R)

    styles = getSampleStyleSheet()

    # ── Color palette ──
    NAVY = HexColor('#1a2744')
    NAVY_LIGHT = HexColor('#2d4a7a')
    STEEL = HexColor('#3d5a80')
    ACCENT = HexColor('#c0392b')
    SUCCESS = HexColor('#27ae60')
    WARNING = HexColor('#e67e22')
    DANGER = HexColor('#c0392b')
    INFO_BLUE = HexColor('#2980b9')
    LIGHT_BG = HexColor('#f4f6f9')
    LIGHTER_BG = HexColor('#fafbfc')
    BORDER = HexColor('#dce1e8')
    TEXT_DARK = HexColor('#1a202c')
    TEXT_BODY = HexColor('#2d3748')
    TEXT_MUTED = HexColor('#718096')
    WHITE = HexColor('#ffffff')

    level_colors = {
        'CRITICAL': DANGER, 'HIGH': WARNING,
        'MEDIUM': INFO_BLUE, 'LOW': SUCCESS, 'INFO': TEXT_MUTED
    }
    color = level_colors.get(level, TEXT_MUTED)

    # ── Typography styles ──
    styles.add(ParagraphStyle('CoverTitle', parent=styles['Title'],
        fontSize=28, textColor=WHITE, alignment=TA_CENTER,
        spaceAfter=6, fontName='Helvetica-Bold', leading=34))
    styles.add(ParagraphStyle('CoverSub', parent=styles['Normal'],
        fontSize=11, textColor=HexColor('#b0bec5'), alignment=TA_CENTER,
        spaceAfter=3, leading=14))
    styles.add(ParagraphStyle('Body', parent=styles['Normal'],
        fontSize=10, leading=15, textColor=TEXT_BODY, spaceBefore=2, spaceAfter=2))
    styles.add(ParagraphStyle('BodySmall', parent=styles['Normal'],
        fontSize=8.5, leading=12, textColor=TEXT_MUTED))
    styles.add(ParagraphStyle('SubHeading', parent=styles['Normal'],
        fontSize=11, textColor=STEEL, spaceBefore=14, spaceAfter=6,
        fontName='Helvetica-Bold', leading=14))
    styles.add(ParagraphStyle('Bullet', parent=styles['Normal'],
        fontSize=10, leading=15, textColor=TEXT_BODY, leftIndent=12,
        spaceBefore=2, spaceAfter=2))
    styles.add(ParagraphStyle('Mono', parent=styles['Normal'],
        fontSize=7.5, fontName='Courier', textColor=HexColor('#4a5568'),
        leading=10, backColor=LIGHTER_BG, spaceBefore=1, spaceAfter=1))
    styles.add(ParagraphStyle('FooterStyle', parent=styles['Normal'],
        fontSize=7, textColor=TEXT_MUTED, alignment=TA_CENTER))

    story = []
    sec = 0  # dynamic section counter

    # ════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════
    story.append(Spacer(1, 55*mm))

    # Title block (navy box)
    cover_content = [
        [Paragraph(
            '<font size="28"><b>RAPPORT D\'INCIDENT</b></font><br/>'
            '<font size="22"><b>PHISHING</b></font>',
            ParagraphStyle('ct', alignment=TA_CENTER, textColor=WHITE, leading=34,
                           fontName='Helvetica-Bold'))]
    ]
    cover_box = Table(cover_content, colWidths=[CONTENT_W], rowHeights=[42*mm])
    cover_box.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), NAVY),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 20),
        ('RIGHTPADDING', (0, 0), (-1, -1), 20),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(cover_box)

    # Severity stripe
    sev_text = f"{level}  —  {classification['priority']}  —  SCORE {score}/100"
    sev_cell = [[Paragraph(
        f'<font color="#ffffff" size="13"><b>{sev_text}</b></font>',
        ParagraphStyle('sev', alignment=TA_CENTER, leading=16))]]
    sev_table = Table(sev_cell, colWidths=[CONTENT_W], rowHeights=[14*mm])
    sev_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), color),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(sev_table)

    story.append(Spacer(1, 12*mm))

    # Cover metadata
    cover_meta = [
        ["Rapport N°", report_number],
        ["Incident", incident_id],
        ["Date", now],
        ["Classification", classification['tlp']],
        ["Categorie", classification['category']],
    ]
    cm_table = Table(cover_meta, colWidths=[45*mm, CONTENT_W - 45*mm])
    cm_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (0, 0), (0, -1), NAVY),
        ('TEXTCOLOR', (1, 0), (1, -1), TEXT_BODY),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LINEBELOW', (0, 0), (-1, -2), 0.3, BORDER),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(cm_table)

    # ── Table of contents ──
    story.append(Spacer(1, 15*mm))
    has_mitre = bool((ai.get("mitre_techniques", []) if ai else []) or
                     (ai.get("social_engineering_tactics", []) if ai else []) or
                     classification.get("mitre_tactics"))
    toc_items = ["Executive Summary", "Alert Details", "Triage Decision", "Enrichment Findings"]
    if has_mitre:
        toc_items.append("MITRE ATT&CK")
    toc_items.append("Recommended Response")

    toc_rows = []
    for i, t in enumerate(toc_items):
        toc_rows.append([f"{i+1}.", t])
    toc_rows.append(["A/B.", "Annexes techniques"])
    toc_t = Table(toc_rows, colWidths=[12*mm, 80*mm])
    toc_t.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (0, 0), (-1, -1), STEEL),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(Paragraph('<b>SOMMAIRE</b>',
        ParagraphStyle('toc_h', fontSize=10, textColor=NAVY, fontName='Helvetica-Bold',
                        spaceAfter=6)))
    story.append(toc_t)

    story.append(PageBreak())

    # ════════════════════════════════════════════════════
    # §1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════
    sec += 1
    story.append(_pdf_section_header_v2(str(sec), "EXECUTIVE SUMMARY", NAVY, CONTENT_W))
    story.append(Spacer(1, 6*mm))

    # Alert box
    one_liner = (
        f"<b>Severity: {level}</b> &mdash; {_safe(classification['category'])} "
        f"detecte depuis <b>{_safe(meta.get('from', 'Inconnu'))}</b> "
        f"(score {score}/100)."
    )
    alert_box = [[Paragraph(one_liner, ParagraphStyle('ab', fontSize=10, leading=14,
                                                       textColor=TEXT_DARK))]]
    ab_t = Table(alert_box, colWidths=[CONTENT_W])
    ab_t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#fff8e1')),
        ('BOX', (0, 0), (-1, -1), 1, HexColor('#ffb300')),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 14),
        ('RIGHTPADDING', (0, 0), (-1, -1), 14),
    ]))
    story.append(ab_t)
    story.append(Spacer(1, 6*mm))

    # AI or verdict summary
    ai_summary = ""
    if ai and not ai.get("error"):
        ai_summary = ai.get("executive_summary", "")
    if ai_summary:
        story.append(Paragraph(_safe(ai_summary), styles['Body']))
        story.append(Spacer(1, 4*mm))
    elif verdict.get("summary"):
        story.append(Paragraph(_safe(verdict["summary"]), styles['Body']))
        story.append(Spacer(1, 4*mm))

    # Impact / Priority / TLP on separate lines for readability
    story.append(Paragraph(
        f"<b>Impact :</b> {_safe(classification['sub_category'])}", styles['Body']))
    story.append(Paragraph(
        f"<b>Priorite :</b> {classification['priority']} &mdash; {classification['severity_description']}", styles['Body']))
    story.append(Paragraph(
        f"<b>Classification TLP :</b> {classification['tlp']}", styles['Body']))

    story.append(PageBreak())

    # ════════════════════════════════════════════════════
    # §2. ALERT DETAILS
    # ════════════════════════════════════════════════════
    sec += 1
    story.append(_pdf_section_header_v2(str(sec), "ALERT DETAILS", NAVY, CONTENT_W))
    story.append(Spacer(1, 6*mm))

    alert_data = [
        ["N° Rapport", report_number],
        ["N° Incident", incident_id],
        ["Date du rapport", now],
        ["Sujet", _safe(meta.get("subject", "N/A"))],
        ["Expediteur", _safe(meta.get("from", "N/A"))],
        ["Destinataire(s)", _safe(meta.get("to", "N/A"))],
        ["Date de reception", _safe(meta.get("date", "N/A"))],
        ["Message-ID", _safe(meta.get("message_id", "N/A"))],
        ["Return-Path", _safe(meta.get("return_path", "N/A"))],
        ["Categorie", classification["category"]],
        ["Score", f"{score}/100 ({level})"],
    ]
    story.append(_pdf_info_table_v2(alert_data, NAVY, CONTENT_W))

    story.append(PageBreak())

    # ════════════════════════════════════════════════════
    # §3. TRIAGE DECISION
    # ════════════════════════════════════════════════════
    sec += 1
    story.append(_pdf_section_header_v2(str(sec), "TRIAGE DECISION", NAVY, CONTENT_W))
    story.append(Spacer(1, 6*mm))

    if score >= 45:
        triage_verdict = "TRUE POSITIVE"
        triage_reason = (
            f"L'analyse a identifie {classification['category'].lower()} avec un score de {score}/100. "
            f"Les indicateurs de compromission confirment une menace reelle."
        )
        triage_color = DANGER
    elif score >= 25:
        triage_verdict = "SUSPICIOUS — INVESTIGATION REQUISE"
        triage_reason = (
            f"Elements suspects detectes (score {score}/100) mais insuffisants pour confirmer. "
            f"Une investigation manuelle est recommandee pour valider la menace."
        )
        triage_color = WARNING
    else:
        triage_verdict = "PROBABLE FALSE POSITIVE"
        triage_reason = (
            f"Score faible ({score}/100). Aucun indicateur de compromission significatif "
            f"n'a ete detecte par l'analyse automatisee."
        )
        triage_color = SUCCESS

    # Verdict banner
    triage_cell = [[Paragraph(
        f'<font color="#ffffff" size="14"><b>{triage_verdict}</b></font>',
        ParagraphStyle('tv', alignment=TA_CENTER, leading=18))]]
    triage_t = Table(triage_cell, colWidths=[CONTENT_W], rowHeights=[16*mm])
    triage_t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), triage_color),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(triage_t)
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph(_safe(triage_reason), styles['Body']))
    story.append(Spacer(1, 8*mm))

    # Scoring breakdown
    factors = risk.get("factors", [])
    if factors:
        story.append(Paragraph("Facteurs de risque", styles['SubHeading']))
        for f_item in factors:
            story.append(Paragraph(f'&bull;  {_safe(f_item)}', styles['Bullet']))
        story.append(Spacer(1, 8*mm))

    # Authentication
    story.append(Paragraph("Authentification email", styles['SubHeading']))
    auth_data = [["PROTOCOLE", "STATUT", "DETAILS"]]
    for proto in ['spf', 'dkim', 'dmarc']:
        item = auth.get(proto, {})
        status = item.get("status", "absent").upper()
        details = _safe(item.get("details", ""))[:70]
        auth_data.append([proto.upper(), status, details])
    t = _pdf_styled_table_v2(auth_data, [30*mm, 28*mm, CONTENT_W - 58*mm], NAVY)
    for i, proto in enumerate(['spf', 'dkim', 'dmarc'], 1):
        status = auth.get(proto, {}).get("status", "absent")
        if status == "pass":
            t.setStyle(TableStyle([('TEXTCOLOR', (1, i), (1, i), SUCCESS)]))
        elif status in ("fail", "softfail"):
            t.setStyle(TableStyle([('TEXTCOLOR', (1, i), (1, i), DANGER)]))
        else:
            t.setStyle(TableStyle([('TEXTCOLOR', (1, i), (1, i), WARNING)]))
    story.append(t)

    # Anomalies
    anomalies = headers_a.get("anomalies", [])
    if anomalies:
        story.append(Spacer(1, 8*mm))
        story.append(Paragraph("Anomalies detectees", styles['SubHeading']))
        for a in anomalies:
            story.append(Paragraph(f'&bull;  {_safe(a)}', styles['Bullet']))

    story.append(PageBreak())

    # ════════════════════════════════════════════════════
    # §4. ENRICHMENT FINDINGS
    # ════════════════════════════════════════════════════
    sec += 1
    story.append(_pdf_section_header_v2(str(sec), "ENRICHMENT FINDINGS", NAVY, CONTENT_W))
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph(
        '<i>IOCs au format defange (hxxps, [.]) pour partage securise.</i>',
        styles['BodySmall']))
    story.append(Spacer(1, 6*mm))

    # Stats bar
    url_count = iocs.get("urls_count", 0)
    suspicious_count = iocs.get("suspicious_urls_count", 0)
    ip_count = iocs.get("ips_count", 0)
    kw_count = iocs.get("keywords_count", 0)
    att_count = len(atts)

    stat_w = CONTENT_W / 5
    stats_data = [[
        f"URLs\n{url_count}", f"Suspectes\n{suspicious_count}",
        f"IPs\n{ip_count}", f"Mots-cles\n{kw_count}", f"PJ\n{att_count}"
    ]]
    st = Table(stats_data, colWidths=[stat_w]*5, rowHeights=[18*mm])
    st.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BACKGROUND', (0, 0), (-1, -1), LIGHT_BG),
        ('TEXTCOLOR', (0, 0), (-1, -1), NAVY),
        ('BOX', (0, 0), (-1, -1), 0.5, BORDER),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, BORDER),
    ]))
    story.append(st)
    story.append(Spacer(1, 8*mm))

    # URLs
    urls = iocs.get("urls", [])
    if urls:
        story.append(Paragraph("URLs detectees", styles['SubHeading']))
        url_data = [["URL (DEFANGED)", "DOMAINE", "FLAGS"]]
        for u in urls[:10]:
            flags = []
            if u.get("suspicious_tld"): flags.append("TLD suspect")
            if u.get("ip_based"): flags.append("IP directe")
            if u.get("url_shortener"): flags.append("Shortener")
            if u.get("mismatched_display"): flags.append("Lien masque")
            url_data.append([
                Paragraph(_safe(_defang_url(u.get("url", "")))[:60],
                          ParagraphStyle('url_cell', fontSize=8.5, leading=11, fontName='Courier')),
                _safe(_defang_domain(u.get("domain", ""))),
                ", ".join(flags) or "—"
            ])
        story.append(_pdf_styled_table_v2(url_data, [85*mm, 42*mm, CONTENT_W - 127*mm], NAVY))
        story.append(Spacer(1, 6*mm))

    # IPs
    ips = iocs.get("ips", [])
    if ips:
        story.append(Paragraph("Adresses IP", styles['SubHeading']))
        ip_data = [["IP (DEFANGED)", "TYPE", "THREAT INTEL"]]
        for ip_info in ips[:10]:
            priv = "Privee" if ip_info.get('is_private') else "Publique"
            ip_str = _defang_ip(ip_info.get("ip", ""))
            ip_data.append([ip_str, priv, f"VT: rechercher {ip_str}"])
        story.append(_pdf_styled_table_v2(ip_data, [42*mm, 28*mm, CONTENT_W - 70*mm], NAVY))
        story.append(Spacer(1, 6*mm))

    # Received chain
    received_chain = headers_a.get("received_chain", [])
    if received_chain:
        story.append(Paragraph("Chaine de transmission (Received hops)", styles['SubHeading']))
        story.append(Paragraph(
            f"L'email a traverse <b>{len(received_chain)}</b> serveur(s).",
            styles['Body']))
        story.append(Spacer(1, 3*mm))
        hop_data = [["HOP", "IP(s)", "DETAILS"]]
        for hop in received_chain:
            ips_str = ", ".join(_defang_ip(ip) for ip in hop.get("ips_found", [])) or "—"
            raw = _safe(hop.get("raw", ""))[:100]
            hop_data.append([
                str(hop.get("hop", "")),
                ips_str,
                Paragraph(raw, ParagraphStyle('hop_d', fontSize=8, leading=10))
            ])
        story.append(_pdf_styled_table_v2(hop_data, [14*mm, 38*mm, CONTENT_W - 52*mm], NAVY))
        story.append(Spacer(1, 6*mm))

    # Phishing keywords
    kws = iocs.get("phishing_keywords", [])
    if kws:
        story.append(Paragraph("Mots-cles phishing detectes", styles['SubHeading']))
        kw_text = ", ".join([f"<b>{_safe(k)}</b>" for k in kws])
        story.append(Paragraph(kw_text, styles['Body']))
        story.append(Spacer(1, 6*mm))

    # Attachments
    if atts:
        story.append(Paragraph("Pieces jointes", styles['SubHeading']))
        for att in atts:
            sus_label = "SUSPECT" if att.get("suspicious_extension") else "Normal"
            sus_color = DANGER if att.get("suspicious_extension") else SUCCESS
            att_info = [
                ["Fichier", _safe(att.get("filename", ""))],
                ["Type MIME", _safe(att.get("content_type", ""))],
                ["Taille", _format_bytes(att.get("size_bytes", 0))],
                ["Statut", sus_label],
                ["MD5", att.get("md5", "N/A")],
                ["SHA256", att.get("sha256", "N/A")],
            ]
            story.append(_pdf_info_table_v2(att_info, sus_color, CONTENT_W))
            sha256 = att.get("sha256", "")
            if sha256 and sha256 != "N/A":
                story.append(Spacer(1, 2*mm))
                story.append(Paragraph(
                    f'<font face="Courier" size="7.5" color="#718096">'
                    f'VirusTotal: hxxps[://]www[.]virustotal[.]com/gui/file/{sha256}</font>',
                    styles['BodySmall']))
            story.append(Spacer(1, 4*mm))

    # ════════════════════════════════════════════════════
    # §5. MITRE ATT&CK MAPPING
    # ════════════════════════════════════════════════════
    mitre = ai.get("mitre_techniques", []) if ai else []
    tactics = ai.get("social_engineering_tactics", []) if ai else []
    if mitre or tactics or classification.get("mitre_tactics"):
        story.append(PageBreak())
        sec += 1
        story.append(_pdf_section_header_v2(str(sec), "MITRE ATT&CK MAPPING", ACCENT, CONTENT_W))
        story.append(Spacer(1, 6*mm))

        if mitre:
            mitre_data = [["ID", "TECHNIQUE", "TACTIC", "PERTINENCE"]]
            for m in mitre:
                mitre_data.append([
                    m.get("id", ""),
                    _safe(m.get("name", "")),
                    _safe(m.get("tactic", "Initial Access")),
                    Paragraph(_safe(m.get("relevance", ""))[:60],
                              ParagraphStyle('mr', fontSize=8.5, leading=11))
                ])
            story.append(_pdf_styled_table_v2(mitre_data, [24*mm, 44*mm, 36*mm, CONTENT_W - 104*mm], ACCENT))
            story.append(Spacer(1, 8*mm))

        if tactics:
            story.append(Paragraph("Tactiques de social engineering (Cialdini)", styles['SubHeading']))
            tac_data = [["TACTIQUE", "PRINCIPE CIALDINI", "EFFICACITE"]]
            for tac in tactics:
                tac_data.append([
                    _safe(tac.get("tactic", "")),
                    _safe(tac.get("cialdini_principle", "")),
                    _safe(tac.get("effectiveness", "")).upper()
                ])
            story.append(_pdf_styled_table_v2(tac_data, [55*mm, 60*mm, CONTENT_W - 115*mm], ACCENT))

        if not mitre and not tactics and classification.get("mitre_tactics"):
            story.append(Paragraph(
                f'Techniques identifiees : {", ".join(classification["mitre_tactics"])}',
                styles['Body']))

    # ════════════════════════════════════════════════════
    # §N. RECOMMENDED RESPONSE
    # ════════════════════════════════════════════════════
    story.append(PageBreak())
    recs = ai.get("recommended_actions", []) if ai else []

    sec += 1
    story.append(_pdf_section_header_v2(str(sec), "RECOMMENDED RESPONSE", NAVY, CONTENT_W))
    story.append(Spacer(1, 6*mm))

    # Baseline recommendations
    if score >= 70:
        baseline_recs = [
            "Bloquer immediatement l'expediteur au niveau du filtre email",
            "Isoler et supprimer l'email de toutes les boites de reception",
            "Verifier les logs d'acces pour detecter des clics sur les liens malveillants",
            "Notifier le RSSI et escalader en incident de securite",
            "Ajouter les IOCs (URLs, IPs, domaines) aux listes de blocage",
        ]
    elif score >= 45:
        baseline_recs = [
            "Mettre l'email en quarantaine pour analyse approfondie",
            "Verifier les logs pour identifier d'autres destinataires",
            "Ajouter les IOCs suspects aux watchlists",
            "Sensibiliser les destinataires au risque de phishing",
        ]
    elif score >= 25:
        baseline_recs = [
            "Surveiller l'expediteur pour activite repetee",
            "Sensibiliser le destinataire aux signaux de phishing",
            "Documenter l'alerte pour reference future",
        ]
    else:
        baseline_recs = [
            "Aucune action immediate requise",
            "Classer comme faux positif si confirme apres verification manuelle",
        ]

    final_recs = recs if recs else baseline_recs
    owners = ["SOC L1", "SOC L2", "SOC L2", "RSSI", "SOC L1"]

    rec_data = [["#", "ACTION", "OWNER"]]
    for i, r in enumerate(final_recs[:6]):
        owner = owners[i] if i < len(owners) else "SOC"
        rec_data.append([str(i + 1), _safe(_clean_recommendation(r)), owner])
    story.append(_pdf_styled_table_v2(rec_data, [12*mm, CONTENT_W - 44*mm, 32*mm], NAVY))

    # ════════════════════════════════════════════════════
    # ANNEXE A — ANALYSE IA DETAILLEE
    # ════════════════════════════════════════════════════
    if ai and not ai.get("error"):
        sem = ai.get("semantic_analysis", {})
        imp = ai.get("impersonation", {})
        soph = ai.get("sophistication", {})

        if sem or imp or soph:
            story.append(PageBreak())
            story.append(_pdf_section_header_v2("A", "ANNEXE — ANALYSE IA DETAILLEE", TEXT_MUTED, CONTENT_W))
            story.append(Spacer(1, 6*mm))

            ai_data = [
                ["Emotion ciblee", _safe(sem.get("target_emotion", "N/A"))],
                ["Pretexte", _safe(sem.get("pretext", imp.get("pretext", "N/A")))],
                ["Entite usurpee", _safe(imp.get("impersonated_entity", "N/A"))],
                ["Qualite usurpation", f"{imp.get('impersonation_quality', 'N/A')}/10"],
                ["Credibilite", f"{sem.get('credibility_assessment', 'N/A')}/10"],
                ["Sophistication", f"{soph.get('level', 'N/A')} ({soph.get('score', 'N/A')}/10)"],
                ["Confiance IA", f"{round((ai.get('ai_confidence', 0)) * 100)}%"],
            ]
            story.append(_pdf_info_table_v2(ai_data, ACCENT, CONTENT_W))

    # ════════════════════════════════════════════════════
    # ANNEXE B — EN-TETES BRUTS
    # ════════════════════════════════════════════════════
    x_headers = headers_a.get("x_headers", {})
    if received_chain or x_headers:
        story.append(PageBreak())
        story.append(_pdf_section_header_v2("B", "ANNEXE — EN-TETES BRUTS (PREUVE)", TEXT_MUTED, CONTENT_W))
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(
            '<i>Conservation de preuve pour analyse independante.</i>',
            styles['BodySmall']))
        story.append(Spacer(1, 6*mm))

        if received_chain:
            story.append(Paragraph("En-tetes Received", styles['SubHeading']))
            for hop in received_chain:
                story.append(Paragraph(
                    f'<font face="Courier" size="7.5" color="#4a5568">'
                    f'Hop {hop.get("hop", "?")}: {_safe(hop.get("raw", ""))}</font>',
                    styles['Mono']))
                story.append(Spacer(1, 2*mm))

        if x_headers:
            story.append(Spacer(1, 6*mm))
            story.append(Paragraph("X-Headers", styles['SubHeading']))
            for key, value in list(x_headers.items())[:20]:
                story.append(Paragraph(
                    f'<font face="Courier" size="7.5" color="#4a5568">'
                    f'{_safe(key)}: {_safe(str(value)[:140])}</font>',
                    styles['Mono']))
                story.append(Spacer(1, 1*mm))

    doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
#
#   DOCX REPORT
#
# ════════════════════════════════════════════════════════════════

def generate_docx_report(analysis: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    NAVY = RGBColor(0x1a, 0x27, 0x44)
    STEEL = RGBColor(0x3d, 0x5a, 0x80)
    PRIMARY = NAVY
    ACCENT = RGBColor(0xc0, 0x39, 0x2b)
    TEXT_DARK = RGBColor(0x1a, 0x20, 0x2c)
    TEXT_BODY = RGBColor(0x2d, 0x37, 0x48)
    TEXT_MUTED = RGBColor(0x71, 0x80, 0x96)
    WHITE = RGBColor(0xff, 0xff, 0xff)
    SUCCESS_RGB = RGBColor(0x27, 0xae, 0x60)
    WARNING_RGB = RGBColor(0xe6, 0x7e, 0x22)
    DANGER_RGB = RGBColor(0xc0, 0x39, 0x2b)

    level_colors = {
        'CRITICAL': DANGER_RGB, 'HIGH': RGBColor(0xe6, 0x7e, 0x22),
        'MEDIUM': RGBColor(0x34, 0x98, 0xdb), 'LOW': SUCCESS_RGB,
        'INFO': RGBColor(0x95, 0xa5, 0xa6)
    }

    meta = analysis.get("metadata", {})
    verdict = analysis.get("verdict", {})
    auth = analysis.get("authentication", {})
    iocs = analysis.get("iocs", {})
    atts = analysis.get("attachments", [])
    risk = analysis.get("risk_scoring", {})
    ai = analysis.get("ai_analysis", {})
    headers_a = analysis.get("headers_analysis", {})

    score = verdict.get("score", 0)
    level = verdict.get("risk_level", "INFO")
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    level_color = level_colors.get(level, TEXT_MUTED)
    incident_id = _generate_incident_id(analysis)
    report_number = _generate_report_number(analysis)
    classification = _classify_incident(analysis)
    timeline = _build_timeline(analysis)
    dsec = 0  # dynamic section counter

    # ── Helpers ──
    def _set_cell_bg(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _section(text, num=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        if num:
            run = p.add_run(f"{num}. ")
            run.font.size = Pt(16)
            run.font.color.rgb = NAVY
            run.bold = True
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        run.bold = True
        border_p = doc.add_paragraph()
        border_p.paragraph_format.space_before = Pt(0)
        border_p.paragraph_format.space_after = Pt(10)
        pPr = border_p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="2" w:color="1A2744"/>'
            f'</w:pBdr>')
        pPr.append(pBdr)

    def _grid_table(data):
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(data):
            for j, val in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = str(val or "")
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.space_before = Pt(4)
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.color.rgb = TEXT_BODY
                if i == 0:
                    _set_cell_bg(cell, "1A2744")
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = WHITE
                            r.bold = True
                elif i % 2 == 0:
                    _set_cell_bg(cell, "F4F6F9")
        return table

    def _kv_table(data):
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, val) in enumerate(data):
            cl = table.rows[i].cells[0]
            cv = table.rows[i].cells[1]
            cl.text = str(label)
            cv.text = str(val or "")
            _set_cell_bg(cl, "F4F6F9")
            for p in cl.paragraphs:
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(4)
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = NAVY
                    r.bold = True
            for p in cv.paragraphs:
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(4)
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = TEXT_BODY
        return table

    def _note(text, italic=True, size=8, color=TEXT_MUTED):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        if italic:
            run.italic = True

    # ════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════
    # Spacer before title
    for _ in range(6):
        doc.add_paragraph()

    # Title block
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    _set_cell_bg(cell, "1A2744")
    cell.height = Cm(5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("RAPPORT D'INCIDENT")
    run.font.size = Pt(28)
    run.font.color.rgb = WHITE
    run.bold = True
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("PHISHING")
    run2.font.size = Pt(22)
    run2.font.color.rgb = WHITE
    run2.bold = True
    doc.add_paragraph()

    # Severity stripe
    sev_t = doc.add_table(rows=1, cols=1)
    sev_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc = sev_t.rows[0].cells[0]
    color_hex = f"{level_color.red:02x}{level_color.green:02x}{level_color.blue:02x}"
    _set_cell_bg(sc, color_hex)
    sc.height = Cm(1.5)
    p = sc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(f"{level}  —  {classification['priority']}  —  SCORE {score}/100")
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    run.bold = True
    doc.add_paragraph()
    doc.add_paragraph()

    # Cover metadata
    cover_meta = [
        ("Rapport N°", report_number),
        ("Incident", incident_id),
        ("Date", now),
        ("Classification", classification['tlp']),
        ("Categorie", classification['category']),
    ]
    _kv_table(cover_meta)

    doc.add_page_break()

    # ════════════════════════════════════════
    # §1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════
    dsec += 1
    _section("EXECUTIVE SUMMARY", str(dsec))

    # One-liner
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    one_liner = (
        f"Severity: {level} — {classification['category']} "
        f"detecte depuis {meta.get('from', 'Inconnu')} (score {score}/100)."
    )
    run = p.add_run(one_liner)
    run.bold = True
    run.font.size = Pt(11)

    # Full summary
    ai_summary = ""
    if ai and not ai.get("error"):
        ai_summary = ai.get("executive_summary", "")
    summary_text = ai_summary or verdict.get("summary", "")
    if summary_text:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(summary_text)
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT_BODY

    # Impact / Priority / TLP on separate lines
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Impact : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(classification['sub_category'])
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Priorite : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(f"{classification['priority']} — {classification['severity_description']}")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Classification TLP : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(classification['tlp'])
    run.font.size = Pt(10)

    # ════════════════════════════════════════
    # 2. ALERT DETAILS
    # ════════════════════════════════════════
    doc.add_page_break()
    dsec += 1
    _section("ALERT DETAILS", str(dsec))
    _kv_table([
        ("N° Rapport", report_number),
        ("N° Incident", incident_id),
        ("Date du rapport", now),
        ("Sujet", meta.get("subject", "N/A")),
        ("Expediteur", meta.get("from", "N/A")),
        ("Destinataire(s)", meta.get("to", "N/A")),
        ("Date reception", meta.get("date", "N/A")),
        ("Message-ID", meta.get("message_id", "N/A")),
        ("Return-Path", meta.get("return_path", "N/A")),
        ("Categorie", classification["category"]),
        ("Score", f"{score}/100 ({level})"),
    ])

    # ════════════════════════════════════════
    # 3. TRIAGE DECISION
    # ════════════════════════════════════════
    doc.add_page_break()
    dsec += 1
    _section("TRIAGE DECISION", str(dsec))

    if score >= 45:
        tv = "TRUE POSITIVE"
        tr = (f"L'analyse a identifie {classification['category'].lower()} avec un score de {score}/100. "
              f"Les indicateurs de compromission confirment une menace reelle.")
        tc = "C0392B"
    elif score >= 25:
        tv = "SUSPICIOUS — INVESTIGATION REQUISE"
        tr = (f"Elements suspects detectes (score {score}/100) mais insuffisants pour confirmer. "
              f"Une investigation manuelle est recommandee pour valider la menace.")
        tc = "E67E22"
    else:
        tv = "PROBABLE FALSE POSITIVE"
        tr = (f"Score faible ({score}/100). Aucun indicateur de compromission significatif "
              f"n'a ete detecte par l'analyse automatisee.")
        tc = "27AE60"

    triage_banner = doc.add_table(rows=1, cols=1)
    triage_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    tcell = triage_banner.rows[0].cells[0]
    _set_cell_bg(tcell, tc)
    tcell.height = Cm(1.5)
    p = tcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(tv)
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(tr)
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_BODY

    # Scoring factors
    factors = risk.get("factors", [])
    if factors:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("Facteurs de risque")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        for f_item in factors:
            bp = doc.add_paragraph(f_item, style='List Bullet')
            for r in bp.runs:
                r.font.size = Pt(10)

    # Auth
    doc.add_paragraph()
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run("Authentification email")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = STEEL

    auth_rows = [["Protocole", "Statut", "Details"]]
    for proto in ['spf', 'dkim', 'dmarc']:
        item = auth.get(proto, {})
        auth_rows.append([proto.upper(), item.get("status", "absent").upper(),
                         (item.get("details", "") or "")[:60]])
    table = _grid_table(auth_rows)
    for i, proto in enumerate(['spf', 'dkim', 'dmarc'], 1):
        status = auth.get(proto, {}).get("status", "absent")
        cell = table.rows[i].cells[1]
        for p in cell.paragraphs:
            for r in p.runs:
                if status == "pass": r.font.color.rgb = SUCCESS_RGB
                elif status in ("fail", "softfail"): r.font.color.rgb = DANGER_RGB
                else: r.font.color.rgb = WARNING_RGB
                r.bold = True

    anomalies = headers_a.get("anomalies", [])
    if anomalies:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("Anomalies detectees")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        for a in anomalies:
            bp = doc.add_paragraph(a, style='List Bullet')
            for r in bp.runs:
                r.font.size = Pt(10)

    # ════════════════════════════════════════
    # 4. ENRICHMENT FINDINGS
    # ════════════════════════════════════════
    doc.add_page_break()
    dsec += 1
    _section("ENRICHMENT FINDINGS", str(dsec))
    _note("IOCs au format defange (hxxps, [.]) pour partage securise.")
    doc.add_paragraph()

    urls = iocs.get("urls", [])
    if urls:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("URLs detectees")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        url_rows = [["URL (defanged)", "Flags"]]
        for u in urls[:10]:
            flags = []
            if u.get("suspicious_tld"): flags.append("TLD suspect")
            if u.get("ip_based"): flags.append("IP directe")
            if u.get("url_shortener"): flags.append("Shortener")
            if u.get("mismatched_display"): flags.append("Lien masque")
            url_rows.append([_defang_url(u.get("url", ""))[:60], ", ".join(flags) or "—"])
        _grid_table(url_rows)
        doc.add_paragraph()

    ips_list = iocs.get("ips", [])
    if ips_list:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("Adresses IP")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        ip_rows = [["IP (defanged)", "Type", "Threat Intel"]]
        for ip_info in ips_list[:10]:
            ip_rows.append([
                _defang_ip(ip_info.get("ip", "")),
                "Privee" if ip_info.get('is_private') else "Publique",
                f"VT: rechercher {_defang_ip(ip_info.get('ip', ''))}"
            ])
        _grid_table(ip_rows)
        doc.add_paragraph()

    received_chain = headers_a.get("received_chain", [])
    if received_chain:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("Chaine de transmission")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        _note(f"L'email a traverse {len(received_chain)} serveur(s).", italic=False, size=10, color=TEXT_BODY)
        hop_rows = [["Hop", "IP(s)", "Details"]]
        for hop in received_chain:
            ips_str = ", ".join(_defang_ip(ip) for ip in hop.get("ips_found", [])) or "—"
            hop_rows.append([str(hop.get("hop", "")), ips_str, hop.get("raw", "")[:100]])
        _grid_table(hop_rows)
        doc.add_paragraph()

    kws = iocs.get("phishing_keywords", [])
    if kws:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("Mots-cles phishing detectes")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        p = doc.add_paragraph()
        run = p.add_run(", ".join(kws))
        run.font.size = Pt(10)
        run.bold = True
        doc.add_paragraph()

    if atts:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run("Pieces jointes")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
        for att in atts:
            _kv_table([
                ("Fichier", att.get("filename", "")),
                ("Type MIME", att.get("content_type", "")),
                ("Taille", _format_bytes(att.get("size_bytes", 0))),
                ("Statut", "SUSPECT" if att.get("suspicious_extension") else "Normal"),
                ("MD5", att.get("md5", "N/A")),
                ("SHA256", att.get("sha256", "N/A")),
            ])
            sha256 = att.get("sha256", "")
            if sha256 and sha256 != "N/A":
                _note(f"VirusTotal: hxxps[://]www[.]virustotal[.]com/gui/file/{sha256}", size=8)
            doc.add_paragraph()

    # ════════════════════════════════════════
    # 5. MITRE ATT&CK MAPPING
    # ════════════════════════════════════════
    mitre = ai.get("mitre_techniques", []) if ai else []
    tactics = ai.get("social_engineering_tactics", []) if ai else []
    if mitre or tactics:
        doc.add_page_break()
        dsec += 1
        _section("MITRE ATT&CK MAPPING", str(dsec))
        if mitre:
            mitre_rows = [["ID", "Technique", "Tactic", "Pertinence"]]
            for m in mitre:
                mitre_rows.append([m.get("id", ""), m.get("name", ""),
                                  m.get("tactic", "Initial Access"), m.get("relevance", "")[:50]])
            _grid_table(mitre_rows)

        if tactics:
            doc.add_paragraph()
            h = doc.add_paragraph()
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run("Tactiques de social engineering (Cialdini)")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = STEEL
            tac_rows = [["Tactique", "Principe Cialdini", "Efficacite"]]
            for tac in tactics:
                tac_rows.append([tac.get("tactic", ""), tac.get("cialdini_principle", ""),
                                tac.get("effectiveness", "").upper()])
            _grid_table(tac_rows)

    # ════════════════════════════════════════
    # N. RECOMMENDED RESPONSE
    # ════════════════════════════════════════
    doc.add_page_break()
    dsec += 1
    _section("RECOMMENDED RESPONSE", str(dsec))

    recs = ai.get("recommended_actions", []) if ai else []
    if score >= 70:
        baseline = [
            "Bloquer immediatement l'expediteur au niveau du filtre email",
            "Isoler et supprimer l'email de toutes les boites de reception",
            "Verifier les logs d'acces pour detecter des clics sur les liens malveillants",
            "Notifier le RSSI et escalader en incident de securite",
            "Ajouter les IOCs (URLs, IPs, domaines) aux listes de blocage",
        ]
    elif score >= 45:
        baseline = [
            "Mettre l'email en quarantaine pour analyse approfondie",
            "Verifier les logs pour identifier d'autres destinataires",
            "Ajouter les IOCs suspects aux watchlists",
            "Sensibiliser les destinataires au risque de phishing",
        ]
    elif score >= 25:
        baseline = [
            "Surveiller l'expediteur pour activite repetee",
            "Sensibiliser le destinataire aux signaux de phishing",
            "Documenter l'alerte pour reference future",
        ]
    else:
        baseline = [
            "Aucune action immediate requise",
            "Classer comme faux positif si confirme apres verification manuelle",
        ]

    final_recs = recs if recs else baseline

    owners = ["SOC L1", "SOC L2", "SOC L2", "RSSI", "SOC L1"]
    rec_rows = [["#", "Action", "Owner"]]
    for i, r in enumerate(final_recs[:6]):
        owner = owners[i] if i < len(owners) else "SOC"
        rec_rows.append([str(i + 1), _clean_recommendation(r), owner])
    _grid_table(rec_rows)

    # ════════════════════════════════════════
    # ANNEXE A — AI DETAILS
    # ════════════════════════════════════════
    if ai and not ai.get("error"):
        sem = ai.get("semantic_analysis", {})
        imp = ai.get("impersonation", {})
        soph = ai.get("sophistication", {})
        if sem or imp or soph:
            doc.add_page_break()
            _section("ANNEXE — ANALYSE IA DETAILLEE", "A")
            _kv_table([
                ("Emotion ciblee", sem.get("target_emotion", "N/A")),
                ("Pretexte", sem.get("pretext", imp.get("pretext", "N/A"))),
                ("Entite usurpee", imp.get("impersonated_entity", "N/A")),
                ("Qualite usurpation", f"{imp.get('impersonation_quality', 'N/A')}/10"),
                ("Credibilite", f"{sem.get('credibility_assessment', 'N/A')}/10"),
                ("Sophistication", f"{soph.get('level', 'N/A')} ({soph.get('score', 'N/A')}/10)"),
                ("Confiance IA", f"{round(ai.get('ai_confidence', 0) * 100)}%"),
            ])

    # ════════════════════════════════════════
    # ANNEXE B — RAW HEADERS
    # ════════════════════════════════════════
    x_headers = headers_a.get("x_headers", {})
    if received_chain or x_headers:
        doc.add_page_break()
        _section("ANNEXE — EN-TETES BRUTS (PREUVE)", "B")
        _note("Conservation de preuve pour analyse independante.")
        doc.add_paragraph()

        if received_chain:
            h = doc.add_paragraph()
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run("En-tetes Received")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = STEEL
            for hop in received_chain:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(f"Hop {hop.get('hop', '?')}: {hop.get('raw', '')}")
                run.font.size = Pt(8)
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

        if x_headers:
            doc.add_paragraph()
            h = doc.add_paragraph()
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run("X-Headers")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = STEEL
            for key, value in list(x_headers.items())[:20]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(f"{key}: {str(value)[:140]}")
                run.font.size = Pt(8)
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    # ════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Rapport {report_number}  |  Incident {incident_id}  |  "
        f"Phishing Email Forensics  |  {now}  |  {classification['tlp']}  |  CONFIDENTIEL")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════
# PDF Helper functions
# ════════════════════════════════════════

def _pdf_section_header(num, title, color):
    """Legacy — kept for backward compatibility."""
    return _pdf_section_header_v2(num, title, color, 166)


def _pdf_section_header_v2(num, title, color, content_w):
    from reportlab.lib.units import mm
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor

    style = ParagraphStyle('sh2', fontSize=14, fontName='Helvetica-Bold',
                           textColor=color, leading=18)
    content = Paragraph(f"{num}. {title}", style)
    t = Table([[content]], colWidths=[content_w], rowHeights=[12*mm])
    t.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LINEBELOW', (0, 0), (-1, -1), 2, color),
    ]))
    return t


def _pdf_styled_table(data, col_widths, header_color):
    """Legacy — kept for backward compatibility."""
    return _pdf_styled_table_v2(data, col_widths, header_color)


def _pdf_styled_table_v2(data, col_widths, header_color):
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle

    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 0), (-1, 0), header_color),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
        ('GRID', (0, 0), (-1, -1), 0.4, HexColor('#dce1e8')),
        ('TOPPADDING', (0, 0), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#ffffff'), HexColor('#f4f6f9')]),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    return t


def _pdf_info_table(data, label_color):
    """Legacy — kept for backward compatibility."""
    return _pdf_info_table_v2(data, label_color, 174)


def _pdf_info_table_v2(data, label_color, content_w):
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle

    label_w = 46*mm
    value_w = content_w - label_w
    t = Table(data, colWidths=[label_w, value_w])
    t.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (0, 0), (0, -1), label_color),
        ('TEXTCOLOR', (1, 0), (1, -1), HexColor('#2d3748')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.4, HexColor('#dce1e8')),
        ('BACKGROUND', (0, 0), (0, -1), HexColor('#f4f6f9')),
        ('TOPPADDING', (0, 0), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    return t
